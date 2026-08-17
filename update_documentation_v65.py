from copy import deepcopy

from docx import Document
from docx.enum.text import WD_COLOR_INDEX


SOURCE = "docs/Caelum_Argenteum_Documentacion_v64_V4.docx"
OUTPUT = "docs/Caelum_Argenteum_Documentacion_v65_V4.docx"

VALIDATED_410 = (
    "Programado y comprobado (4.10): Actor.Inv incorpora un lingote de hierro "
    "apilable, una llave de plata Key no apilable y una carta sellada como "
    "objeto clave único, todos con peso base 0,1. Las pruebas confirmaron los "
    "ocho filtros, el peso, las pilas y el traslado de materiales y objetos "
    "clave a la Caja Mágica. LOCKDEFS 200 exige la llave de plata; las llaves "
    "permanecen en el inventario personal para que la comprobación sea nativa."
)

CATALOGUE_RULES = [
    "Catálogo 4.11 — piezas principales: hoja, hoja pequeña, hoja curva, hoja larga, hoja ancha, asta, armazón, armazón largo, cabeza de arma, cabeza redonda, placa, placa redonda, placa de lágrima, placa de torre, placa mágica, placa grande y cota de malla.",
    "Catálogo 4.11 — recursos: tejido, cuero, esencias de fuego, agua, tierra y viento, y quintaesencia. Componentes secundarios: empuñadura, empuñadura larga, punta, mango, mango largo, cuerda, cuerda reforzada, correa, correa reforzada, cañón, mecanismo y bases de bastón, campana, libro y estatuilla.",
    "Tiers de material: metal = bronce/hierro/acero; madera = común/dura/ébano o mágica; esencia = simple/fina/pura; cuero = vaca/cocodrilo o tiburón/demonio o dragón; tejido = lana/algodón/seda. Los componentes secundarios genéricos no cambian por tier.",
    "Pilas: tipo y tier forman la identidad de la pila. Dos unidades solo se apilan si ambos coinciden; tiers distintos permanecen como objetos separados. Cada unidad conserva peso base 0,1 y toda la pila ocupa un espacio de Caja Mágica.",
    "Prueba de llave sin mapa: el comando ca_debug_test_silver_lock ejecuta la comprobación nativa de LOCKDEFS 200. Sin la llave muestra el mensaje de bloqueo; con la llave de plata confirma el acceso. Una puerta real debe declarar el número 200 en su especial de línea.",
]

IMPLEMENTATION_411 = (
    "Programado para probar (4.11): el inventario nativo ofrece el catálogo "
    "parametrizado de materiales, selección de tipo y tier, pilas separadas por "
    "identidad, peso, desvío por sobrecarga, Caja Mágica y descarte. El comando "
    "ca_debug_test_silver_lock comprueba directamente LOCKDEFS 200 sin editar "
    "un mapa."
)


def replace_and_highlight(paragraph, text):
    run_properties = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        run_properties = deepcopy(paragraph.runs[0]._r.rPr)
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(text)
    if run_properties is not None:
        if run._r.rPr is not None:
            run._r.remove(run._r.rPr)
        run._r.insert(0, run_properties)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def insert_after(document, anchor, text, style_name):
    paragraph = document.add_paragraph(style=style_name)
    run = paragraph.add_run(text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    anchor._p.addnext(paragraph._p)
    return paragraph


def replace_cell_and_highlight(cell, text):
    paragraph = cell.paragraphs[0]
    replace_and_highlight(paragraph, text)


document = Document(SOURCE)
catalogue_anchor = None
implementation_anchor = None

for paragraph in document.paragraphs:
    if paragraph.text == (
        'Contador: Muestra la cantidad total de cada material '
        '(ej. "Hoja de bronce: 3").'
    ):
        catalogue_anchor = paragraph
    if paragraph.text.startswith("Programado para probar (4.10):"):
        replace_and_highlight(paragraph, VALIDATED_410)
        implementation_anchor = paragraph

if catalogue_anchor is None or implementation_anchor is None:
    raise RuntimeError("No se encontraron los anclajes documentales 4.11")

anchor = catalogue_anchor
for rule in CATALOGUE_RULES:
    anchor = insert_after(document, anchor, rule, "List Bullet")

insert_after(document, implementation_anchor, IMPLEMENTATION_411, "List Bullet")

# La tabla de recetas conservaba dos pesos anteriores a la regla definitiva.
shield_table = document.tables[41]
if shield_table.cell(2, 2).text != "14" or shield_table.cell(3, 2).text != "18":
    raise RuntimeError("La tabla de escudos no coincide con la versión v64")
replace_cell_and_highlight(shield_table.cell(2, 2), "12")
replace_cell_and_highlight(shield_table.cell(3, 2), "16")

document.core_properties.title = "Caelum Argenteum — Documentación v65 / Versión 4"
document.core_properties.subject = (
    "Catálogo nativo de materiales y prueba LOCKDEFS 200 — 4.11"
)
document.save(OUTPUT)
