from copy import deepcopy

from docx import Document
from docx.enum.text import WD_COLOR_INDEX


SOURCE = "docs/Caelum_Argenteum_Documentacion_v65_V4.docx"
OUTPUT = "docs/Caelum_Argenteum_Documentacion_v66_V4.docx"


def replace_paragraph(paragraph, text, highlight=True):
    run_properties = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        run_properties = deepcopy(paragraph.runs[0]._r.rPr)
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(text)
    if run_properties is not None:
        if run._r.rPr is not None:
            run._r.remove(run._r.rPr)
        run._r.insert(0, run_properties)
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def replace_cell(cell, text, highlight=True):
    paragraph = cell.paragraphs[0]
    replace_paragraph(paragraph, text, highlight)


def insert_after(document, anchor, text, style_name="List Bullet"):
    paragraph = document.add_paragraph(style=style_name)
    run = paragraph.add_run(text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    anchor._p.addnext(paragraph._p)
    return paragraph


def rebuild_table(table, rows):
    template = deepcopy(table.rows[1]._tr)
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)
    for values, highlight in rows:
        new_tr = deepcopy(template)
        table._tbl.append(new_tr)
        row = table.rows[-1]
        for cell, value in zip(row.cells, values):
            replace_cell(cell, value, highlight)


document = Document(SOURCE)

# Resumen general del crafteo físico.
for paragraph in document.paragraphs:
    if paragraph.text.startswith("Crafteo de armas físicas ("):
        replace_paragraph(
            paragraph,
            "Crafteo de armas físicas mediante dos componentes: uno principal "
            "y uno secundario. Las familias definitivas incluyen daga, hachuela, "
            "machete, jabalina, espada, hacha, mangual, lanza, espadón, hacha de "
            "guerra, alabarda, guanteletes gigantes, arco estándar, carabina, "
            "arco largo y ballesta. El componente indicado por cada receta "
            "determina el tier.",
        )
    elif paragraph.text.startswith("El RetrocesoBase es el empuje del ataque"):
        replace_paragraph(
            paragraph,
            "El RetrocesoBase es el empuje propio del ataque (por ejemplo, 100 "
            "para un golpe contundente pesado).",
        )
    elif paragraph.text.startswith("En combate: Recibe un golpe de maza"):
        replace_paragraph(
            paragraph,
            paragraph.text.replace("golpe de maza", "golpe contundente"),
        )

# Materiales base: elimina componentes descartados y agrega los definitivos.
material_table = document.tables[33]
replace_cell(
    material_table.cell(1, 1),
    "Hoja, Hoja pequeña, Hoja curva, Hoja larga, Hoja ancha, Asta, Punta, "
    "Cabeza de arma, Cabeza pequeña, Cabeza redonda, Mango, Mango largo, "
    "Empuñadura, Empuñadura larga, Cadena, Placa grande y Correa reforzada",
)
replace_cell(
    material_table.cell(1, 2),
    "Crear las doce armas cuerpo a cuerpo definitivas de las familias 2, 3 y 4.",
)
replace_cell(
    material_table.cell(2, 1),
    "Armazón, Armazón largo, Cuerda, Cuerda reforzada, Cañón y Mecanismo",
)
replace_cell(
    material_table.cell(2, 2),
    "Crear arco estándar, carabina, arco largo y ballesta.",
)

recipes = [
    (("Hoja pequeña + Empuñadura", "Daga", "2", "Hoja pequeña", "Más rápida, menos daño."), False),
    (("Cabeza pequeña + Mango", "Hachuela", "2", "Cabeza pequeña", "El tier lo aporta la cabeza pequeña."), True),
    (("Hoja curva + Mango", "Machete", "2", "Hoja curva", "El tier lo aporta la hoja curva."), True),
    (("Asta + Punta", "Jabalina", "2", "Punta", "Comparte componentes con la lanza; puede arrojarse."), True),
    (("Hoja + Empuñadura", "Espada", "3", "Hoja", "Estándar."), False),
    (("Cabeza de arma + Mango", "Hacha", "3", "Cabeza de arma", "Daño medio-alto."), True),
    (("Cabeza redonda + Cadena", "Mangual", "3", "Cabeza redonda", "La cadena es genérica y no posee tier."), True),
    (("Asta + Punta", "Lanza", "3", "Punta", "Con escudo puede atacar mientras bloquea."), True),
    (("Hoja larga + Empuñadura larga", "Espadón", "4", "Hoja larga", "Dos manos, muy pesado."), False),
    (("Hoja ancha + Mango largo", "Hacha de guerra", "4", "Hoja ancha", "Dos manos, mucho daño."), True),
    (("Hoja larga + Asta", "Alabarda", "4", "Hoja larga", "Mucho alcance, dos manos."), False),
    (("Placa grande + Correa reforzada", "Guanteletes gigantes", "4", "Placa grande", "Dos manos, bloqueo sin escudo."), False),
    (("Armazón + Cuerda", "Arco estándar", "5", "Armazón", "Equilibrado."), False),
    (("Cañón + Mecanismo", "Carabina", "5", "Cañón", "Gran daño y alcance; disparo lento y dispersión extrema."), False),
    (("Armazón largo + Cuerda", "Arco largo", "5", "Armazón largo", "Preciso, largo alcance, caro."), True),
    (("Armazón + Cuerda reforzada", "Ballesta", "5", "Armazón", "Punto medio entre arco estándar y largo."), False),
    (("Placa de escudo + Correa", "Escudo", "-", "Placa", "El tipo depende de la placa: redonda, lágrima, torre o mágica."), True),
]
rebuild_table(document.tables[39], recipes)

disassembly = [
    (("Daga", "Hoja pequeña", "Empuñadura", "Se conserva el tier de la hoja pequeña."), False),
    (("Hachuela", "Cabeza pequeña", "Mango", "Se conserva el tier de la cabeza pequeña."), True),
    (("Machete", "Hoja curva", "Mango", "Se conserva el tier de la hoja curva."), True),
    (("Jabalina", "Punta", "Asta", "Se conserva el tier de la punta."), True),
    (("Espada", "Hoja", "Empuñadura", "Se conserva el tier de la hoja."), False),
    (("Hacha", "Cabeza de arma", "Mango", "Se conserva el tier de la cabeza."), True),
    (("Mangual", "Cabeza redonda", "Cadena", "Se conserva el tier de la cabeza redonda."), True),
    (("Lanza", "Punta", "Asta", "Se conserva el tier de la punta."), True),
    (("Espadón", "Hoja larga", "Empuñadura larga", "Se conserva el tier de la hoja larga."), False),
    (("Hacha de guerra", "Hoja ancha", "Mango largo", "Se conserva el tier de la hoja ancha."), True),
    (("Alabarda", "Hoja larga", "Asta", "Se conserva el tier de la hoja larga."), False),
    (("Guanteletes gigantes", "Placa grande", "Correa reforzada", "Se conserva el tier de la placa grande."), False),
    (("Arco estándar", "Armazón", "Cuerda", "Se conserva el tier del armazón."), False),
    (("Carabina", "Cañón", "Mecanismo", "Se conserva el tier del cañón."), True),
    (("Arco largo", "Armazón largo", "Cuerda", "Se conserva el tier del armazón largo."), True),
    (("Ballesta", "Armazón", "Cuerda reforzada", "Se conserva el tier del armazón."), True),
    (("Escudo", "Placa", "Correa", "Se conserva el tier de la placa."), False),
    (("Arma de esencia", "Esencia u objeto base", "El otro componente", "Se elige cuál de los dos componentes recuperar."), True),
    (("Armadura", "Pieza principal", "Correa", "Se conserva el tier de la pieza."), False),
]
rebuild_table(document.tables[45], disassembly)

# Estaciones y ejemplos de daño ya no deben nombrar armas eliminadas.
replace_cell(
    document.tables[46].cell(1, 2),
    "Armas físicas (espadas, dagas, hachas, hachuelas, manguales, lanzas, "
    "jabalinas, escudos, espadones, etc.) y armaduras metálicas.",
)
replace_cell(
    document.tables[86].cell(1, 3),
    "Mangual y guanteletes gigantes.",
)
replace_cell(
    document.tables[122].cell(1, 2),
    "Espadón (cortante), mangual (contundente), guanteletes gigantes (contundente).",
)

push_table = document.tables[123]
for row in list(push_table.rows)[1:]:
    if row.cells[0].text.startswith("Maza"):
        push_table._tbl.remove(row._tr)
        break

# Catálogo activo y estado de implementación 4.12.
catalogue_anchor = None
implementation_anchor = None
for paragraph in document.paragraphs:
    if paragraph.text.startswith("Catálogo 4.11 — piezas principales:"):
        replace_paragraph(
            paragraph,
            "Catálogo activo 4.12 — piezas principales: hoja, hoja pequeña, "
            "hoja curva, hoja larga, hoja ancha, asta, armazón, armazón largo, "
            "cabeza de arma, cabeza pequeña, cabeza redonda, placa, placa "
            "redonda, placa de lágrima, placa de torre, placa mágica, placa "
            "grande y cota de malla.",
        )
        catalogue_anchor = paragraph
    elif paragraph.text.startswith("Catálogo 4.11 — recursos:"):
        replace_paragraph(
            paragraph,
            paragraph.text.replace(
                "mecanismo y bases", "mecanismo, cadena y bases"
            ),
        )
    elif paragraph.text.startswith("Programado para probar (4.11):"):
        implementation_anchor = paragraph

if catalogue_anchor is None or implementation_anchor is None:
    raise RuntimeError("No se encontraron los anclajes documentales 4.12")

insert_after(
    document,
    implementation_anchor,
    "Programado para probar (4.12): el código incorpora el catálogo autoritativo "
    "de las 16 armas físicas, con familia, daños primario/secundario, tipos de "
    "daño, velocidad, alcance, dispersión, crítico, aire y recetas. El tier de "
    "lanza y jabalina lo determina la punta; el de mangual, la cabeza redonda. "
    "La cadena es genérica. Los 41 materiales activos poseen al menos una "
    "receta; el antiguo lingote de prueba queda oculto y se conserva solo para "
    "compatibilidad. El consumo de cantidades se implementará cuando se defina "
    "su fórmula definitiva. El comando ca_debug_audit_crafting_catalogue debe "
    "informar 16 recetas, 41 materiales activos y 0 sin uso.",
)

document.core_properties.title = "Caelum Argenteum — Documentación v66 / Versión 4"
document.core_properties.subject = (
    "Catálogo definitivo de armas físicas y recetas estructurales — 4.12"
)
document.save(OUTPUT)
