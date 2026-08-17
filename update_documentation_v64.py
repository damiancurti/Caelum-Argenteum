from copy import deepcopy

from docx import Document
from docx.enum.text import WD_COLOR_INDEX


SOURCE = "docs/Caelum_Argenteum_Documentacion_v63_V4.docx"
OUTPUT = "docs/Caelum_Argenteum_Documentacion_v64_V4.docx"


REPLACEMENTS = {
    "Los materiales conservan su bolsillo separado y no ocupan inventario personal ni espacios en la Caja Mágica.":
        "Los materiales son objetos nativos apilables del inventario personal. Pesan 0,1 por unidad de forma predeterminada y una pila completa puede ocupar un solo espacio de Caja Mágica, donde aporta peso 0.",
    "Caja Mágica y materiales: La caja recibe excedentes de peso y tiene espacios limitados; los materiales usan un bolsillo separado.":
        "Caja Mágica: Recibe excedentes de peso y tiene espacios limitados. Materiales y objetos clave siguen perteneciendo al inventario nativo y pueden guardarse allí; las llaves permanecen en el inventario personal para conservar la comprobación nativa de cerraduras.",
    "Filtros: Todos, armas, armaduras, escudos, accesorios y otros.":
        "Etiquetas y filtros: armas, armaduras y escudos —equipados o desequipados—, consumibles, munición, materiales, llaves, objetos clave y estado «en Caja Mágica». Cada categoría se muestra en una sección distinta.",
    "Los materiales permanecen fuera de ambos conteos y conservan su pestaña propia.":
        "Materiales, llaves y objetos clave permanecen en Actor.Inv y cuentan en el inventario y en el peso. Conservan pestañas propias para encontrarlos sin mezclarlos con equipo, consumibles o munición.",
    "2.3. Sección de materiales (subpestaña)":
        "2.3. Secciones de materiales, llaves y objetos clave",
    "Una pestaña separada dentro del inventario dedicada exclusivamente a los materiales de crafteo.":
        "Tres secciones separadas dentro del inventario agrupan materiales de crafteo, llaves de acceso y objetos clave de misión o interacción.",
    "No ocupan espacio en la caja mágica (están en un \"bolsillo de materiales\" separado).":
        "Los materiales son apilables; cada unidad pesa 0,1 por defecto. Una pila en Caja Mágica pesa 0 y ocupa un solo espacio sin importar la cantidad.",
    "Si la Caja Mágica está llena y el objeto no cabe por peso, no puede recogerse. Cada objeto no apilable ocupa un espacio; toda una pila apilable ocupa uno solo. Los materiales no usan ninguno de estos espacios.":
        "Si la Caja Mágica está llena y el objeto no cabe por peso, no puede recogerse. Cada objeto no apilable ocupa un espacio y toda una pila apilable ocupa uno solo. Materiales y objetos clave aplican esta regla; las llaves nativas no se envían a la Caja Mágica.",
    "Materiales: No ocupan espacio en la caja. Se almacenan en una sección separada y siempre están disponibles para el crafteo.":
        "Materiales: Son pilas nativas con peso unitario y sección propia. Pueden almacenarse en la Caja Mágica; allí la pila ocupa un espacio, pesa 0 y continúa disponible como propiedad del jugador.",
    "Programado para probar (4.9): los cinco consumibles son pilas nativas con peso real, desvío por sobrecarga, un único espacio de Caja Mágica por pila, uso nativo y Powerups de 10 segundos. El menú compacto permite crearlos en el suelo, usarlos, guardarlos, recuperarlos y soltarlos.":
        "Programado y comprobado (4.9): los cinco consumibles son pilas nativas con peso real, desvío por sobrecarga, un único espacio de Caja Mágica por pila, uso nativo y Powerups de 10 segundos. Las pruebas confirmaron creación, recogida, uso, almacenamiento, recuperación y descarte.",
}

KEY_RULES = [
    "Llaves: Derivan de Key, no son apilables y pesan 0,1 por defecto. LOCKDEFS permite exigirlas en puertas o en acciones mediante Door_LockedRaise, Door_Animated, Generic_Door, ACS_LockedExecute y ACS_LockedExecuteDoor.",
    "Caja Mágica y llaves: Una llave permanece siempre en el inventario personal. LOCKDEFS comprueba posesión nativa y no reconoce una bandera interna de Caja Mágica; esta restricción evita que una llave guardada abra igualmente una cerradura.",
    "Objetos clave: Son instancias únicas, no apilables y pesan 0,1 por defecto. Pueden entrar en la Caja Mágica. Los objetos de puzzle consumibles podrán derivar de PuzzleItem cuando una interacción deba retirarlos al usarse.",
    "Infraestructura de misiones: GZDoom aporta llaves, PuzzleItem, QuestItem, diálogos de Strife, mensajes y UI programable. Caelum reutilizará esas piezas, pero mantendrá una capa propia para objetivos, seguimiento y presentación de misiones.",
]

IMPLEMENTATION_NOTE = (
    "Programado para probar (4.10): Actor.Inv incorpora un lingote de hierro "
    "apilable, una llave de plata Key no apilable y una carta sellada como "
    "objeto clave único, todos con peso base 0,1. El inventario compacto posee "
    "ocho filtros y LOCKDEFS 200 permite exigir la llave de plata en puertas o "
    "acciones bloqueadas. Materiales y objetos clave pueden alternar Caja "
    "Mágica; las llaves permanecen en el inventario personal."
)


def replace_and_highlight(paragraph, text):
    run_properties = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        run_properties = deepcopy(paragraph.runs[0]._r.rPr)
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(text)
    if run_properties is not None:
        if run._r.rPr is not None:
            run._r.remove(run._r.rPr)
        run._r.insert(0, run_properties)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def insert_after(document, anchor, text, style_name):
    paragraph = document.add_paragraph(style=style_name)
    run = paragraph.add_run(text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    anchor._p.addnext(paragraph._p)
    return paragraph


document = Document(SOURCE)
found = set()
category_anchor = None
implementation_anchor = None

for paragraph in list(document.paragraphs):
    original = paragraph.text
    replacement = REPLACEMENTS.get(original)
    if replacement is None:
        continue
    replace_and_highlight(paragraph, replacement)
    found.add(original)
    if original.startswith("Los materiales permanecen fuera de ambos conteos"):
        category_anchor = paragraph
    if original.startswith("Programado para probar (4.9)"):
        implementation_anchor = paragraph

missing = set(REPLACEMENTS) - found
if missing:
    raise RuntimeError(f"No se encontraron {len(missing)} pasajes: {sorted(missing)}")
if category_anchor is None or implementation_anchor is None:
    raise RuntimeError("No se encontraron los puntos de inserción 4.10")

anchor = category_anchor
for rule in KEY_RULES:
    anchor = insert_after(document, anchor, rule, "List Bullet")

insert_after(document, implementation_anchor, IMPLEMENTATION_NOTE, "List Bullet")

document.core_properties.title = "Caelum Argenteum — Documentación v64 / Versión 4"
document.core_properties.subject = (
    "Inventario categorizado, materiales, llaves y objetos clave 4.10"
)
document.save(OUTPUT)
